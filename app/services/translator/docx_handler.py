import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from typing import Optional, List, Tuple, Dict
from .core import HYMTTranslator
import zipfile
from pathlib import Path
from .chunk_processor import ChunkProcessor

class DocxHandler:
    """Word 文档 (.docx) 处理器（支持：原顺序+合并单元格+嵌套表格）"""

    def __init__(self, translator: HYMTTranslator):
        self.translator = translator
        self.chunk_processor = ChunkProcessor(translator.model_name)

    def _process_single_table_to_html(self, table, target_lang, show_bilingual) -> str:
        """
        处理单个表格转HTML（支持合并单元格+嵌套表格）
        :param table: docx Table 对象
        :param target_lang: 目标语言
        :param show_bilingual: 是否显示中英对照
        :return: 带合并单元格+嵌套表格的表格HTML
        """
        # 解析表格的合并单元格信息
        merge_info = self._parse_table_merge_info(table)
        html_rows = []

        # 逐行处理表格
        for row_idx, row in enumerate(table.rows):
            html_cells = []
            for col_idx, cell in enumerate(row.cells):
                # 跳过被合并的单元格（已被其他单元格覆盖）
                if (row_idx, col_idx) in merge_info['merged_cells']:
                    continue

                # 获取单元格合并属性
                colspan = merge_info['colspan'].get((row_idx, col_idx), 1)
                rowspan = merge_info['rowspan'].get((row_idx, col_idx), 1)

                # 核心修改：解析单元格内的内容（文本 + 嵌套表格）
                cell_content_html = self._parse_cell_content(cell, target_lang, show_bilingual)

                # 生成带合并属性的单元格HTML
                cell_attrs = []
                if colspan > 1:
                    cell_attrs.append(f'colspan="{colspan}"')
                if rowspan > 1:
                    cell_attrs.append(f'rowspan="{rowspan}"')
                attrs_str = ' '.join(cell_attrs) if cell_attrs else ''

                cell_html = f'<td {attrs_str}>{cell_content_html}</td>' if attrs_str else f'<td>{cell_content_html}</td>'
                html_cells.append(cell_html)

            # 生成行HTML
            html_rows.append(f'<tr>{"".join(html_cells)}</tr>')

        # 生成完整表格HTML（优化样式适配嵌套表格）
        table_html = f"""
        <div class="document-table" style="margin: 20px 0; overflow-x: auto;">
            <table border="1" cellpadding="8" cellspacing="0" style="border-collapse: collapse; width: 100%;">
                {"".join(html_rows)}
            </table>
        </div>
        """
        return table_html

    def _parse_cell_content(self, cell, target_lang, show_bilingual) -> str:
        """
        解析单元格内容（文本 + 嵌套表格）
        :param cell: docx Cell 对象
        :param target_lang: 目标语言
        :param show_bilingual: 是否显示中英对照
        :return: 单元格内容的 HTML 字符串
        """
        content_parts = []

        # 1. 提取单元格内的文本
        cell_text = cell.text.strip()
        if cell_text:
            # 检测是否为中文，如果是则跳过翻译
            if self._is_chinese_text(cell_text):
                # 中文不翻译，原文和译文相同
                text_html = f"""
                <div class="cell-text">
                    <div class="original-text">{cell_text}</div>
                </div>
                """
                content_parts.append(text_html)
            elif show_bilingual:
                translated_text = self.translator.translate_text(cell_text, target_lang)
                text_html = f"""
                <div class="cell-text">
                    <div class="original-text">{cell_text}</div>
                    <div class="translated-text">{translated_text}</div>
                </div>
                """
                content_parts.append(text_html)
            else:
                translated_text = self.translator.translate_text(cell_text, target_lang)
                text_html = f'<div class="cell-text translated-text">{translated_text}</div>'
                content_parts.append(text_html)

        # 2. 检查并解析单元格内的嵌套表格（核心新增逻辑）
        # 遍历单元格的 XML 节点，查找嵌套的<table>节点
        cell_elem = cell._tc
        nested_tables = cell_elem.xpath('.//w:tbl')
        if nested_tables:
            # 遍历所有嵌套表格，递归处理
            for nested_tbl_elem in nested_tables:
                # 从 XML 节点构建 docx Table 对象
                from docx.table import Table
                nested_table = Table(nested_tbl_elem, cell)
                # 递归生成嵌套表格的 HTML
                nested_table_html = self._process_single_table_to_html(nested_table, target_lang, show_bilingual)
                # 添加嵌套表格样式（缩进 + 边框区分）
                styled_nested_table = f"""
                <div class="nested-table" style="margin: 10px 0 10px 20px; border: 1px dashed #666; padding: 5px;">
                    {nested_table_html}
                </div>
                """
                content_parts.append(styled_nested_table)

        # 3. 处理空单元格
        if not content_parts:
            return '<div class="empty-cell">&nbsp;</div>'

        return ''.join(content_parts)


    def _parse_table_merge_info(self, table) -> Dict[str, any]:
        """
        解析表格的合并单元格信息（保留原有逻辑）
        """
        colspan = {}
        rowspan = {}
        merged_cells = set()

        # 遍历所有单元格，解析合并属性
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 跳过已标记为合并的单元格
                if (row_idx, col_idx) in merged_cells:
                    continue

                # 解析横向合并（gridSpan）
                cell_elem = cell._tc
                grid_span = cell_elem.xpath('.//w:gridSpan')
                if grid_span:
                    span = int(grid_span[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
                    colspan[(row_idx, col_idx)] = span
                    # 标记被合并的列
                    for c in range(1, span):
                        merged_cells.add((row_idx, col_idx + c))

                # 解析纵向合并（vMerge）
                v_merge = cell_elem.xpath('.//w:vMerge')
                if v_merge:
                    v_merge_val = v_merge[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                    # vMerge="restart" 表示合并起始单元格，需要计算合并行数
                    if v_merge_val == 'restart':
                        merge_rows = 1
                        # 向下遍历，统计连续的合并行
                        for r in range(row_idx + 1, len(table.rows)):
                            next_cell = table.rows[r].cells[col_idx]
                            next_v_merge = next_cell._tc.xpath('.//w:vMerge')
                            if next_v_merge and next_v_merge[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') is None:
                                merge_rows += 1
                                merged_cells.add((r, col_idx))
                            else:
                                break
                        rowspan[(row_idx, col_idx)] = merge_rows

        return {
            'colspan': colspan,
            'rowspan': rowspan,
            'merged_cells': merged_cells
        }

    # 保留原有 convert_to_html/_parse_document_elements/_process_single_paragraph_to_html 等核心方法
    def convert_to_html(
            self,
            input_path: str,
            output_dir: str,
            target_lang: str = "Chinese",
            show_bilingual: bool = True,
            translate_all: int = 0,
            preserve_original_styles: bool = True
    ) -> str:
        """
        将 Word 文档转换为 HTML 并翻译（支持：原顺序 + 合并单元格 + 嵌套表格）
        """
        os.makedirs(output_dir, exist_ok=True)
        print(f"Processing Word to HTML (full flow): {input_path}")
        doc = Document(input_path)

        # 【新增】文档复杂度检测
        complexity = self._detect_document_complexity(doc)
        use_batch_translation = complexity['recommend_batch']

        # 构建 HTML 内容（增强样式适配）
        html_content = [self._get_html_header(preserve_original_styles)]

        # 进度追踪
        tracker = self.translator.get_progress_tracker()
        # 解析文档所有原生元素（段落/表格/图片）并计数
        all_elements = self._parse_document_elements(doc, input_path)
        total_elements = len(all_elements)
        elements_to_process = total_elements if translate_all == 0 else min(translate_all, total_elements)
        tracker.set_file_info(os.path.basename(input_path), elements_to_process, "element")

        # 核心：按原生顺序处理每个元素
        processed_count = 0
        img_idx = 0  # 图片索引（用于命名）

        # 修改：为每个文档创建独立的图片文件夹
        base_name = os.path.basename(input_path)
        name_without_ext = os.path.splitext(base_name)[0]
        img_folder_name = f"{name_without_ext}_images"
        img_dir = os.path.join(output_dir, img_folder_name)
        os.makedirs(img_dir, exist_ok=True)

        # 【关键修改】根据文档复杂度选择翻译策略
        translated_results = None

        if use_batch_translation:
            # 【批量翻译模式】提前批量翻译所有段落
            if translate_all == 0:
                paras_to_process = sum(1 for elem_type, _ in all_elements if elem_type == "paragraph")
            else:
                paras_to_process = min(translate_all,
                                       sum(1 for elem_type, _ in all_elements if elem_type == "paragraph"))

            # 收集所有需要翻译的段落
            paragraphs_to_translate = []
            para_idx = 0
            for elem_type, elem_data in all_elements:
                if elem_type == "paragraph" and para_idx < paras_to_process:
                    para = elem_data
                    text = para.text.strip()
                    style = self._get_paragraph_style_type(para)
                    paragraphs_to_translate.append((text, style))
                    para_idx += 1

            # 执行批量翻译
            if paragraphs_to_translate:
                print(f"\n[DOCX 转换] 使用批量翻译模式，共 {len(paragraphs_to_translate)} 个段落...")
                translated_results = self._batch_translate_paragraphs(
                    paragraphs_to_translate,
                    target_lang,
                    show_bilingual
                )
        else:
            # 【逐段翻译模式】不提前翻译，在处理时逐段翻译
            print(f"\n[DOCX 转换] 检测到复杂表格，使用逐段翻译模式以保证质量...")
            print(f"[总体进度] 文档共 {elements_to_process} 个元素（段落 + 表格 + 图片），开始处理...\n")

        # 重置段落索引，用于从批量翻译结果中获取
        para_result_idx = 0

        for elem_idx, (elem_type, elem_data) in enumerate(all_elements):
            if translate_all > 0 and processed_count >= elements_to_process:
                break

            # 【新增】计算并显示总体进度（每个元素都显示）
            current_progress = (elem_idx + 1) / len(all_elements) * 100
            progress_bar = self._create_progress_bar(current_progress, width=30)

            if elem_type == "paragraph":
                para = elem_data
                # 从批量翻译结果中获取译文（如果使用了批量翻译）
                if use_batch_translation and translated_results and para_result_idx < len(translated_results):
                    orig, trans, style = translated_results[para_result_idx]

                    # 根据 show_bilingual 生成 HTML
                    if not orig.strip():
                        html = '<br/>'
                    elif self._is_chinese_text(orig):
                        # 中文段落不翻译
                        html = self._generate_paragraph_html(orig, orig, style, preserve_original_styles)
                    elif show_bilingual:
                        html = self._generate_paragraph_html(orig, trans, style, preserve_original_styles)
                    else:
                        html = self._generate_paragraph_html("", trans, style, preserve_original_styles)

                    para_result_idx += 1
                else:
                    # 【修改】未使用批量翻译或超出范围，使用逐段翻译并显示进度
                    # 逐段翻译
                    text = para.text.strip()
                    if not text:
                        html = '<br/>'
                        print(
                            f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} [空段落]",
                            end="", flush=True)
                    elif self._is_chinese_text(text):
                        # 中文段落不翻译
                        para_style = self._get_paragraph_style_type(para)
                        html = self._generate_paragraph_html(text, text, para_style, preserve_original_styles)
                        print(
                            f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} [中文跳过]",
                            end="", flush=True)
                    else:
                        # 需要翻译的段落
                        try:
                            translated = self.translator.translate_text(text, target_lang)
                            para_style = self._get_paragraph_style_type(para)

                            if show_bilingual:
                                html = self._generate_paragraph_html(text, translated, para_style,
                                                                     preserve_original_styles)
                            else:
                                html = self._generate_paragraph_html("", translated, para_style,
                                                                     preserve_original_styles)

                            print(
                                f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} ✓ {len(translated)}字",
                                end="", flush=True)
                        except Exception as e:
                            fallback_text = f"[翻译失败：{str(e)}]"
                            para_style = self._get_paragraph_style_type(para)
                            html = self._generate_paragraph_html("", fallback_text, para_style,
                                                                 preserve_original_styles)
                            print(
                                f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} ✗ {fallback_text}",
                                end="", flush=True)

                if html:
                    html_content.append(html)
                processed_count += 1
                tracker.update_paragraph(processed_count)

            elif elem_type == "table":
                table = elem_data
                html = self._process_single_table_to_html(table, target_lang, show_bilingual)
                html_content.append(html)
                print(f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} [表格]",
                      end="", flush=True)
                processed_count += 1
                tracker.update_paragraph(processed_count)

            elif elem_type == "image":
                # 处理单张图片（按出现顺序）
                img_file = elem_data
                img_ext = Path(img_file).suffix
                img_name = f"image_{img_idx}{img_ext}"
                img_save_path = os.path.join(img_dir, img_name)
                # 提取图片文件
                with zipfile.ZipFile(input_path) as doc_zip:
                    with open(img_save_path, 'wb') as f:
                        f.write(doc_zip.read(img_file))
                # 生成图片 HTML（使用独立文件夹路径）
                img_rel_path = os.path.join(img_folder_name, img_name)
                img_html = f'''
                        <div class="document-image">
                            <img src="{self._escape_html(img_rel_path)}" alt="文档图片 {img_idx + 1}" 
                                 style="max-width: 100%; height: auto; margin: 20px 0;">
                        </div>
                        '''
                html_content.append(img_html)
                print(f"\r[{progress_bar}] {current_progress:.1f}% | 元素 {elem_idx + 1}/{len(all_elements)} [图片]",
                      end="", flush=True)
                img_idx += 1
                processed_count += 1
                tracker.update_paragraph(processed_count)

        # 换行，避免覆盖最后的进度信息
        print(f"\n[完成] 文档处理完毕，共处理 {processed_count} 个元素")

        # 闭合 HTML 标签
        html_content.append(self._get_html_footer())

        # 保存 HTML 文件
        output_path = os.path.join(output_dir, f"{name_without_ext}_translated.html")

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(html_content))

        tracker.mark_completed()
        print(f"翻译后的 HTML（完整样式 + 原顺序 + 合并单元格 + 嵌套表格）已保存至：{output_path}")
        return output_path

    def process(
            self,
            input_path: str,
            output_path: Optional[str] = None,
            target_lang: str = "Chinese",
            translate_all: int = 0
    ) -> str:
        """
        处理 DOCX 文档翻译（智能选择批量/逐段翻译）
        :param input_path: 输入文件路径
        :param output_path: 输出文件路径（可选）
        :param target_lang: 目标语言
        :param translate_all: 是否翻译全文，0=全文，>0 表示翻译前 N 个段落
        :return: 输出文件路径
        """
        if not output_path:
            base, ext = os.path.splitext(input_path)
            output_path = f"{base}_translated.txt"

        print(f"Processing DOCX: {input_path}")
        doc = Document(input_path)

        # 【新增】文档复杂度检测
        complexity = self._detect_document_complexity(doc)
        use_batch_translation = complexity['recommend_batch']

        # 收集所有段落文本和样式
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            style = self._get_paragraph_style_type(para)
            if text:  # 只保留有内容的段落
                paragraphs.append((text, style))

        tracker = self.translator.get_progress_tracker()
        paras_to_process = len(paragraphs) if translate_all == 0 else min(translate_all, len(paragraphs))
        tracker.set_file_info(os.path.basename(input_path), paras_to_process, "paragraph")

        # 【关键修改】根据文档复杂度选择翻译策略
        if use_batch_translation:
            # 批量翻译模式
            print(f"\n[DOCX 处理] 使用批量翻译模式，共 {len(paragraphs[:paras_to_process])} 个段落...")
            translated_results = self._batch_translate_paragraphs(
                paragraphs[:paras_to_process],
                target_lang,
                show_bilingual=False
            )
        else:
            # 逐段翻译模式
            print(f"\n[DOCX 处理] 检测到复杂表格，使用逐段翻译模式...")
            translated_results = []
            for idx, (text, style) in enumerate(paragraphs[:paras_to_process]):
                if not text.strip():
                    translated_results.append((text, text, style))
                elif self._is_chinese_text(text):
                    print(f"[跳过] 段落 {idx + 1} 为中文，已跳过翻译")
                    translated_results.append((text, text, style))
                else:
                    try:
                        translated_para = self.translator.translate_text(text, target_lang)
                        translated_results.append((text, translated_para, style))
                        print(f"  ✓ 段落 {idx + 1}: {len(translated_para)} 字")
                        tracker.update_paragraph(idx + 1)
                    except Exception as e:
                        fallback_text = f"[翻译失败：{str(e)}]"
                        translated_results.append((text, fallback_text, style))
                        print(f"  ✗ 段落 {idx + 1}: {fallback_text}")
                        tracker.update_paragraph(idx + 1)

        # 未翻译的段落保持原样
        final_results = translated_results + paragraphs[paras_to_process:]

        # 生成输出（保持原有格式）
        results = []
        for idx, ((orig, _), (_, trans, _)) in enumerate(zip(paragraphs, final_results)):
            if orig.strip():
                results.append(f"{orig}\n\n{trans}\n\n{'-' * 30}\n\n")

        with open(output_path, 'w', encoding='utf-8') as f:
            f.write("".join(results))

        tracker.mark_completed()
        print(f"TXT saved to: {output_path}")
        return output_path


    def _parse_document_elements(self, doc: Document, doc_path: str) -> List[Tuple[str, object]]:
        """解析DOCX文档的所有元素，按原生顺序返回"""
        elements = []
        # 1. 先解析段落和表格（基于docx库的节点遍历）
        for child in doc.element.body:
            tag = child.tag.rpartition('}')[-1]  # 获取节点标签（p/tbl）
            if tag == 'p':
                # 找到对应的Paragraph对象
                para_idx = len([e for e in elements if e[0] == "paragraph"])
                elements.append(("paragraph", doc.paragraphs[para_idx]))
            elif tag == 'tbl':
                # 找到对应的Table对象
                table_idx = len([e for e in elements if e[0] == "table"])
                elements.append(("table", doc.tables[table_idx]))

        # 2. 解析图片（关联到段落位置）
        # 先提取所有图片文件
        doc_zip = zipfile.ZipFile(doc_path)
        image_files = [f for f in doc_zip.namelist() if f.startswith('word/media/')]
        doc_zip.close()

        # 遍历段落，找到包含图片的段落，插入到对应位置
        img_idx = 0
        new_elements = []
        for elem in elements:
            new_elements.append(elem)
            # 检查当前段落是否包含图片
            if elem[0] == "paragraph":
                para = elem[1]
                if para._element.xpath('.//w:drawing') and img_idx < len(image_files):
                    new_elements.append(("image", image_files[img_idx]))
                    img_idx += 1

        return new_elements

    def _process_single_paragraph_to_html(self, para, target_lang, show_bilingual, preserve_styles):
        """处理单个段落转 HTML（复用原有逻辑）"""
        text = para.text.strip()
        if not text:
            return '<br/>'

        # 识别段落样式
        para_style = self._get_paragraph_style_type(para)

        # 检测是否为中文，如果是则跳过翻译
        if self._is_chinese_text(text):
            # 中文段落不翻译，直接返回原文
            return self._generate_paragraph_html(text, text, para_style, preserve_styles)

        if show_bilingual:
            translated = self.translator.translate_text(text, target_lang)
            return self._generate_paragraph_html(text, translated, para_style, preserve_styles)
        else:
            translated = self.translator.translate_text(text, target_lang)
            return self._generate_paragraph_html("", translated, para_style, preserve_styles)

    def _batch_translate_paragraphs(
            self,
            paragraphs: list,
            target_lang: str,
            show_bilingual: bool
    ) -> list:
        """
        批量翻译段落（利用大模型上下文窗口）
        :param paragraphs: 段落列表（每个元素为 (text, para_style) 元组）
        :param target_lang: 目标语言
        :param show_bilingual: 是否显示中英对照
        :return: 翻译后的段落列表（每个元素为 (original, translated) 元组）
        """
        if not paragraphs:
            return []

        # 1. 过滤掉中文段落（不翻译）
        processed_paragraphs = []
        translation_needed = []

        for idx, (text, style) in enumerate(paragraphs):
            if not text.strip():
                processed_paragraphs.append((text, text, style))
            elif self._is_chinese_text(text):
                print(f"[跳过] 段落 {idx + 1} 为中文，已跳过翻译")
                processed_paragraphs.append((text, text, style))
            else:
                processed_paragraphs.append(None)  # 占位，稍后填充
                translation_needed.append((idx, text, style))

        if not translation_needed:
            return processed_paragraphs

        # 2. 创建带映射的分块
        chunks = self.chunk_processor.create_chunks_with_mapping(
            [text for _, text, _ in translation_needed],
            target_lang
        )

        print(f"\n[批量翻译] 共 {len(translation_needed)} 段需要翻译，分为 {len(chunks)} 个批次")

        # 3. 逐批翻译
        translated_idx = 0
        for chunk_idx, chunk in enumerate(chunks):
            # 显示进度条
            current_progress = (chunk_idx + 1) / len(chunks) * 100
            progress_bar = self._create_progress_bar(current_progress, width=30)
            print(f"\r[{progress_bar}] {current_progress:.1f}% | 批次 {chunk_idx + 1}/{len(chunks)}", end="",
                  flush=True)

            try:
                # 调用翻译（一次性翻译整个 chunk）
                translated_chunk_text = self.translator.translate_text(
                    chunk["text"],
                    target_lang
                )

                # 解析翻译结果，还原为段落列表
                translated_paras = self.chunk_processor.parse_translated_chunks(
                    translated_chunk_text,
                    len(chunk["paragraph_indices"])
                )

                # 将翻译结果精确映射回原位置
                for para_local_idx, global_para_idx in enumerate(chunk["paragraph_indices"]):
                    original_idx, original_text, style = translation_needed[global_para_idx]

                    if para_local_idx < len(translated_paras):
                        translated_para = translated_paras[para_local_idx]
                        processed_paragraphs[original_idx] = (original_text, translated_para, style)
                        print(f"  ✓ 段落 {original_idx + 1}: {len(translated_para)} 字")
                    else:
                        # 段落数量不匹配时的容错
                        fallback_text = f"[部分翻译失败：期望{len(chunk['paragraph_indices'])}段，实际{len(translated_paras)}段]"
                        processed_paragraphs[original_idx] = (original_text, fallback_text, style)
                        print(f"  ✗ 段落 {original_idx + 1}: {fallback_text}")

                    translated_idx += 1

            except Exception as e:
                print(f"\n[错误] 批次 {chunk_idx + 1} 翻译失败：{e}")
                # 失败回退：逐段翻译
                for para_local_idx, global_para_idx in enumerate(chunk["paragraph_indices"]):
                    original_idx, original_text, style = translation_needed[global_para_idx]

                    try:
                        translated_para = self.translator.translate_text(original_text, target_lang)
                        processed_paragraphs[original_idx] = (original_text, translated_para, style)
                        print(f"  ✓ 段落 {original_idx + 1} (回退): {len(translated_para)} 字")
                    except Exception as e2:
                        fallback_text = f"[翻译失败：{str(e2)}]"
                        processed_paragraphs[original_idx] = (original_text, fallback_text, style)
                        print(f"  ✗ 段落 {original_idx + 1} (回退): {fallback_text}")

                    translated_idx += 1

        print("\n")  # 换行
        return processed_paragraphs

    def _is_chinese_text(self, text: str) -> bool:
        """
        检测文本是否主要为中文（包括阿拉伯数字和常见标点）
        :param text: 待检测文本
        :return: True 表示主要是中文、数字或标点，False 表示非中文
        """
        import re
        # 统计中文字符、阿拉伯数字和中文标点的数量
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', text)
        arabic_numerals = re.findall(r'[0-9]', text)
        # 中文标点符号范围：包括常用标点如，。！？；：""''（）【】《》等
        chinese_punctuation = re.findall(r'[\u3000-\u303f\uff00-\uffef]', text)

        # 如果中文字符、阿拉伯数字或中文标点占比超过 70%，则认为是中文文本
        if len(text) > 0:
            chinese_or_digit_count = len(chinese_chars) + len(arabic_numerals) + len(chinese_punctuation)
            ratio = chinese_or_digit_count / len(text)
            return ratio >= 0.7
        return False

    def _create_progress_bar(self, percentage: float, width: int = 30) -> str:
        """创建进度条字符串"""
        percentage = max(0, min(100, percentage))
        filled_length = int(width * percentage / 100)
        bar = '█' * filled_length + '░' * (width - filled_length)
        return bar

# 辅助方法（优化样式适配嵌套表格）
    def _get_html_header(self, preserve_styles: bool) -> str:
        """生成HTML头部（补充嵌套表格样式）"""
        style = """
        <style>
            body { font-family: Arial, sans-serif; line-height: 1.6; padding: 20px; }
            .document-table table { border: 1px solid #000; }
            .document-table td { border: 1px solid #000; vertical-align: top; padding: 8px; }
            /* 嵌套表格样式 */
            .nested-table { border: 1px dashed #666 !important; margin: 10px; padding: 5px; }
            .nested-table .document-table { margin: 0 !important; }
            /* 单元格文本样式 */
            .cell-text { margin-bottom: 5px; }
            .original-text { color: #333; margin-bottom: 3px; }
            .translated-text { color: #0066cc; }
            .empty-cell { color: #ccc; }
            .document-image { text-align: center; }
            h1, h2, h3 { margin: 15px 0; }
            ul, ol { margin: 10px 0 10px 20px; }
        </style>
        """ if preserve_styles else ""
        return f"""
        <!DOCTYPE html>
        <html lang="zh-CN">
        <head>
            <meta charset="UTF-8">
            <title>Translated Document</title>
            {style}
        </head>
        <body>
        """

    def _get_html_footer(self) -> str:
        """生成HTML尾部"""
        return """
        </body>
        </html>
        """

    def _escape_html(self, text: str) -> str:
        """转义HTML特殊字符"""
        return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    def _get_paragraph_style_type(self, para) -> str:
        """识别段落样式类型（标题/列表/普通文本）"""
        style_name = para.style.name.lower()
        if 'heading 1' in style_name or '标题1' in style_name:
            return 'h1'
        elif 'heading 2' in style_name or '标题2' in style_name:
            return 'h2'
        elif 'heading 3' in style_name or '标题3' in style_name:
            return 'h3'
        elif para.style.name in ['List Paragraph', '列表段落']:
            return 'list'
        else:
            return 'p'

    def _generate_paragraph_html(self, original: str, translated: str, style_type: str, preserve_styles: bool) -> str:
        """生成段落HTML（保留样式）"""
        if not preserve_styles:
            style_type = 'p'

        # 【关键修复】如果原文和译文相同（如中文段落），只输出一行，避免重复
        if original == translated:
            content = f'<span class="original-text">{original}</span>'
        else:
            content = ""
            if original:
                content += f'<span class="original-text">{original}</span><br/>'
            content += f'<span class="translated-text">{translated}</span>'

        if style_type == 'h1':
            return f'<h1>{content}</h1>'
        elif style_type == 'h2':
            return f'<h2>{content}</h2>'
        elif style_type == 'h3':
            return f'<h3>{content}</h3>'
        elif style_type == 'list':
            return f'<li>{content}</li>'
        else:
            return f'<p>{content}</p>'


    def _has_complex_tables(self, doc) -> bool:
        """
        检测文档是否包含复杂表格（合并单元格、嵌套表格等）
        :param doc: Document 对象
        :return: True 表示包含复杂表格
        """
        for table in doc.tables:
            # 检查是否有合并单元格
            merge_info = self._parse_table_merge_info(table)
            if merge_info['colspan'] or merge_info['rowspan']:
                print(f"[复杂表格检测] 发现合并单元格表格")
                return True

            # 检查是否有嵌套表格
            for row in table.rows:
                for cell in row.cells:
                    cell_elem = cell._tc
                    nested_tables = cell_elem.xpath('.//w:tbl')
                    if nested_tables:
                        print(f"[复杂表格检测] 发现嵌套表格")
                        return True

        return False


    def _detect_document_complexity(self, doc) -> dict:
        """
        检测文档复杂度
        :param doc: Document 对象
        :return: 复杂度检测结果
        """
        complexity = {
            'has_complex_tables': False,
            'total_paragraphs': 0,
            'total_tables': len(doc.tables),
            'recommend_batch': True
        }

        # 统计段落数
        complexity['total_paragraphs'] = sum(1 for para in doc.paragraphs if para.text.strip())

        # 检测表格复杂度
        complexity['has_complex_tables'] = self._has_complex_tables(doc)

        # 决策：如果有复杂表格，建议使用逐段翻译
        if complexity['has_complex_tables']:
            complexity['recommend_batch'] = False
            print(
                f"[文档复杂度检测] 文档包含 {complexity['total_paragraphs']} 个段落，{complexity['total_tables']} 个表格（含复杂表格）")
            print(f"[文档复杂度检测] 建议使用逐段翻译模式以保证表格处理质量")
        else:
            print(
                f"[文档复杂度检测] 文档包含 {complexity['total_paragraphs']} 个段落，{complexity['total_tables']} 个简单表格")
            print(f"[文档复杂度检测] 建议使用批量翻译模式以提升速度")

        return complexity

